#!/usr/bin/env python3
"""Generate a reference.docx for pandoc with bordered Table Grid style as default.

Pandoc's markdown→docx by default applies the "Table" style (no borders). To
get visible borders on every table, we tell pandoc to use this reference doc
via `--reference-doc=reference.docx`. The reference's Normal Table is set to
match the built-in "Table Grid" pattern (1pt single black borders on all
sides + inside).

Run once; reuse for all OnE renders.
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_table_grid_default(docx: Document) -> None:
    """Override the default Table style so every pandoc-emitted table has borders."""
    styles_xml = docx.styles.element

    # Find or create the default Table style
    # Pandoc uses style name "Table" for emitted tables; we'll add borders to it.
    for style_el in styles_xml.findall(qn("w:style")):
        if style_el.get(qn("w:styleId")) in ("Table", "TableGrid", "Tablegrid"):
            # Drop existing tblBorders if any
            existing = style_el.find(qn("w:tblPr"))
            if existing is not None:
                style_el.remove(existing)
            tblPr = OxmlElement("w:tblPr")
            tblBorders = OxmlElement("w:tblBorders")
            for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
                b = OxmlElement(f"w:{side}")
                b.set(qn("w:val"), "single")
                b.set(qn("w:sz"), "4")        # 0.5pt
                b.set(qn("w:space"), "0")
                b.set(qn("w:color"), "888888") # medium grey
                tblBorders.append(b)
            tblPr.append(tblBorders)
            style_el.append(tblPr)

    # Also add an explicit "Table" style entry if pandoc emits w:tblStyle val="Table"
    have_table = any(
        s.get(qn("w:styleId")) == "Table"
        for s in styles_xml.findall(qn("w:style"))
    )
    if not have_table:
        new_style = OxmlElement("w:style")
        new_style.set(qn("w:type"), "table")
        new_style.set(qn("w:styleId"), "Table")
        name = OxmlElement("w:name"); name.set(qn("w:val"), "Table")
        new_style.append(name)
        basedOn = OxmlElement("w:basedOn"); basedOn.set(qn("w:val"), "TableNormal")
        new_style.append(basedOn)
        tblPr = OxmlElement("w:tblPr")
        tblBorders = OxmlElement("w:tblBorders")
        for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
            b = OxmlElement(f"w:{side}")
            b.set(qn("w:val"), "single")
            b.set(qn("w:sz"), "4")
            b.set(qn("w:space"), "0")
            b.set(qn("w:color"), "888888")
            tblBorders.append(b)
        tblPr.append(tblBorders)
        new_style.append(tblPr)
        styles_xml.append(new_style)


def main() -> None:
    out = Path("/Users/ag/Desktop/AIProjectsJuly2025/TIUConsulting/10X Door/CA properties/titlePro/docs/Cure_Response/reference.docx")
    doc = Document()
    # Insert a placeholder paragraph + table so pandoc can copy the styles.
    doc.add_paragraph("placeholder")
    t = doc.add_table(rows=2, cols=2)
    t.style = "Table Grid"
    for r in t.rows:
        for c in r.cells:
            c.text = "x"
    set_table_grid_default(doc)
    doc.save(str(out))
    print(f"[+] Wrote {out} ({out.stat().st_size:,} bytes)")


if __name__ == "__main__":
    main()
